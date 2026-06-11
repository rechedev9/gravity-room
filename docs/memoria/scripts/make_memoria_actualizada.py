import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont

ROOT = Path('/home/reche/projects/TrackerRSN')
text = (ROOT/'body_clean.txt').read_text(encoding='utf-8')
text = text.replace('\u200b', '')

summary = '''1. Resumen
Este proyecto presenta el diseño, desarrollo y despliegue de Gravity Room, una aplicación web full-stack para el seguimiento de programas de entrenamiento de fuerza progresivos. La hipótesis de partida es que muchos programas de fuerza pueden representarse como una combinación de días, ejercicios, etapas y reglas de progresión; por tanto, el cálculo del siguiente entrenamiento puede automatizarse mediante un motor común, en lugar de programar una lógica distinta para cada rutina.

Frente a la mayoría de aplicaciones del sector —que funcionan principalmente como registros pasivos de sesiones de gimnasio—, Gravity Room calcula qué peso, series y repeticiones corresponden a la siguiente sesión a partir del historial del usuario y de las reglas del programa. El objetivo práctico es reducir la carga mental del deportista: el usuario registra lo que ha completado y la aplicación decide el siguiente paso de forma coherente y auditable.

La arquitectura se compone de un cliente SPA desarrollado con React 19 y TanStack Router/Query, instalable como PWA; una API REST tipada sobre Bun y ElysiaJS; persistencia en PostgreSQL mediante Drizzle ORM; autenticación con Google OAuth y rotación de JWT; y despliegue en un VPS propio de Hetzner Cloud mediante Docker Compose, Caddy como reverse proxy con TLS automático e integración continua desde GitHub Actions. El motor de progresión está escrito en TypeScript y se comparte entre cliente y servidor, garantizando que ambos calculan los mismos resultados.

Como conclusión, el proyecto demuestra que es viable construir un producto full-stack completo, con un modelo de dominio no trivial, desplegado públicamente y operable por terceros, con un coste de infraestructura reducido —alrededor de 13,50 € mensuales— y manteniendo buenas prácticas profesionales: tipado extremo a extremo, pruebas automatizadas, observabilidad estructurada, migraciones automáticas y trazabilidad por commit. La separación entre definición del programa y estado del usuario permite añadir nuevos programas al catálogo principalmente como configuración, no como programación específica, validando así la hipótesis inicial del proyecto.

Enlaces del proyecto
Aplicación en vivo: https://gravityroom.app
Repositorio (AGPL-3.0): https://github.com/rechedev9/gravity-room
'''

index = '''2. Índice
1. Resumen
2. Índice
3. Justificación
3.1 ¿Por qué este proyecto?
3.2 ¿Para qué?
3.3 Estado del arte y alternativas
3.4 Relevancia para el ciclo formativo
3.5 Origen de la idea
4. Objetivos
4.1 Objetivo general
4.2 Objetivos específicos
5. Desarrollo
5.1 Contextualización del sector profesional
5.2 Desarrollo del proyecto
5.2.1 Elección del stack y justificación
5.2.2 Arquitectura general
5.2.3 Motor de progresión declarativo
5.2.4 Despliegue e integración continua
5.2.5 Modelo de datos
5.2.6 API REST y autenticación
5.2.7 Aplicación web (SPA + PWA + modo invitado)
5.2.8 Diseño de interfaz y experiencia de usuario
5.2.9 Control de versiones con Git y GitHub
5.2.10 Calidad y verificación
5.2.11 Hitos de desarrollo
6. Conclusiones
6.1 Cumplimiento de los objetivos
6.2 Aprendizajes
6.3 Líneas futuras
6.4 Reflexión final
7. Bibliografía
8. Anexos
8.1 Anexo I. Glosario de términos técnicos
8.2 Anexo II. Definición declarativa reducida del programa
'''

stack = '''5.2.1 Elección del stack y justificación
La selección de tecnologías responde a tres criterios ordenados: (1) compartir código entre cliente y servidor siempre que sea posible, para garantizar consistencia del motor de progresión; (2) favorecer el tipado extremo a extremo, para detectar errores antes de ejecutar la aplicación; y (3) minimizar la complejidad operacional, para poder desplegar en un VPS modesto sin depender de plataformas propietarias.

Bun + TypeScript en backend y herramientas. Bun es un runtime de JavaScript que integra gestor de paquetes, ejecutor de tests, bundler y servidor HTTP nativo (Bun Team, 2025). En Gravity Room la motivación principal no es únicamente el rendimiento, sino la homogeneidad: usar TypeScript en cliente, servidor y paquete compartido permite que el motor de progresión viva en un único sitio y se ejecute igual en ambas capas. Esto evita duplicar reglas como add_weight, advance_stage o deload_percent, reduciendo el riesgo de que la API y la SPA calculen progresiones distintas.

ElysiaJS para la API REST. La API se implementa con ElysiaJS, un framework orientado a Bun que permite declarar rutas, validaciones y tipos de forma compacta (ElysiaJS Team, 2025). Esta elección encaja con el proyecto porque la API no necesita una arquitectura empresarial pesada: expone autenticación, catálogo, instancias de programas, resultados y deshacer. Elysia permite mantener esos endpoints tipados sin introducir una capa excesiva de abstracción.

React 19 + Vite + TanStack Router/Query. React 19 se utiliza como base de la SPA por su ecosistema y por las mejoras de compilación y renderizado introducidas en la versión actual (Meta / React Team, 2024). React Compiler reduce parte del código defensivo necesario para memorizar componentes y hooks (React Compiler Team, 2025). TanStack Router aporta rutas tipadas y TanStack Query gestiona caché, reintentos, invalidación y actualizaciones optimistas (TanStack, 2025). En Gravity Room esto es relevante porque el usuario interactúa constantemente con datos remotos —programas, sesiones, progreso, resultados— y necesita una interfaz rápida incluso cuando la red no responde de forma inmediata.

PostgreSQL + Drizzle ORM. PostgreSQL aporta consistencia transaccional, integridad referencial y soporte para JSONB (PostgreSQL Global Development Group, 2025). En este proyecto esa consistencia es importante porque completar, fallar o deshacer una serie modifica el historial que después utiliza el motor de progresión. Drizzle ORM se ha elegido porque define el esquema desde TypeScript, genera migraciones SQL versionadas y conserva el tipado entre base de datos y aplicación (Drizzle Team, 2025). Los campos JSONB se usan sólo donde aportan flexibilidad real, como definiciones de programas o detalle de series, y no como sustituto de un modelo relacional.

Tailwind CSS 4 para la interfaz. Tailwind CSS permite construir la interfaz mediante clases utilitarias y mantener un sistema visual consistente sin diseñar componentes desde cero en cada pantalla (Tailwind Labs, 2025). En Gravity Room resulta útil porque la aplicación necesita muchas vistas móviles con una misma estética —tema oscuro, acento ámbar, tarjetas de programa, botones grandes durante el entrenamiento— y cambios rápidos durante el desarrollo.

Hetzner Cloud + Docker Compose + Caddy. Hetzner Cloud ofrece VPS en Europa con buen rendimiento por coste y firewall gestionado a nivel de red (Hetzner Online GmbH, 2026). Docker Compose es suficiente para un sistema pequeño formado por SPA estática, API, PostgreSQL y Redis, evitando la complejidad de Kubernetes para un proyecto que no la necesita (Docker Inc., 2025). Caddy actúa como reverse proxy y automatiza los certificados HTTPS mediante Let's Encrypt (Caddy Authors, 2025). En la práctica, esta combinación permite desplegar gravityroom.app con bajo coste, TLS automático y una configuración comprensible para un único desarrollador.

Zod en frontend y backend. Zod permite definir esquemas de validación reutilizables en cliente y servidor. En Gravity Room esto evita que el navegador acepte una estructura de datos que luego la API rechace, o al revés. Los mismos esquemas validan formularios, importación del modo invitado y campos JSONB persistidos en PostgreSQL, cerrando una clase frecuente de errores entre capas.
'''

motor = '''5.2.3 Motor de progresión declarativo
El corazón del proyecto es el motor de progresión computeGenericProgram(), implementado en TypeScript en el paquete compartido @gzclp/shared. Dicho de forma sencilla, este motor es la parte de la aplicación que decide qué debe hacer el usuario en el próximo entrenamiento. Para calcularlo utiliza tres elementos: el programa elegido, los pesos de referencia del usuario y el historial de sesiones completadas.

La idea de diseño es separar definición y estado. La definición describe el programa de forma estable: qué días tiene, qué ejercicios aparecen en cada día, qué etapas existen y qué regla se aplica cuando el usuario acierta o falla. El estado, en cambio, representa la situación actual del usuario: qué peso toca hoy, en qué etapa está cada ejercicio y cuál será el siguiente entrenamiento. Gravity Room no guarda ese estado como un valor manual que pueda quedarse desactualizado; lo recalcula a partir del historial. Este enfoque es similar a un event sourcing simplificado, pero aplicado a un dominio pequeño y controlado.

Este diseño tiene tres ventajas prácticas. Primero, reduce errores de sincronización: si dos dispositivos tienen el mismo historial, ambos reconstruyen el mismo estado. Segundo, permite añadir programas nuevos como configuración, no como código específico para cada rutina. Tercero, permite ejecutar el motor en el navegador, lo que hace posible el modo invitado sin conexión ni cuenta.

Las reglas de progresión disponibles son seis: add_weight, que incrementa el peso; advance_stage, que pasa a la siguiente etapa de series y repeticiones; advance_stage_add_weight, que combina ambas acciones; deload_percent, que reduce el peso un porcentaje para facilitar la recuperación; add_weight_reset_stage, que sube peso y vuelve a la primera etapa; y no_change, que mantiene la prescripción actual. Estas reglas permiten modelar programas conocidos como GZCLP, StrongLifts 5x5 o variantes de 5/3/1. En el caso de GZCLP, la estructura de progresión se inspira en la descripción pública del programa (Lexa, 2019).

Modelo declarativo en TypeScript
Una definición de programa es un objeto que enumera días, ejercicios y reglas. El fragmento siguiente muestra una versión reducida de la estructura. La versión completa puede consultarse en el Anexo II.

Figura 9. Extracto reducido de una definición declarativa de programa en TypeScript. Elaboración propia.

```typescript
export interface ProgramDefinition {
  id: string;         // "gzclp", "stronglifts-5x5", ...
  name: string;
  level: "beginner" | "intermediate" | "advanced";
  days: Day[];
}

export interface Day {
  label: string;      // "Día A", "Día B", ...
  slots: Slot[];      // ejercicios del día
}

export interface Slot {
  exerciseId: string;
  stages: Stage[];   // etapas progresivas del ejercicio
}

export interface Stage {
  sets: number;
  reps: number;
  amrap: boolean;
  successRule: ProgressionRule;
  failureRule: ProgressionRule;
}
```

El punto importante es que ProgressionRule es la única parte que el motor interpreta de forma especial. Todo lo demás —días, ejercicios, etapas o intensidades— es contenido configurable. Por eso añadir un programa nuevo al catálogo no exige modificar el algoritmo central, sino publicar una definición que respete el contrato.

Caso trabajado: GZCLP, sentadilla T1
Para ilustrar el funcionamiento, conviene seguir un ejemplo real. En GZCLP, un ejercicio principal como la sentadilla T1 puede pasar por varias etapas: 5×3+, 6×2+ y 10×1+. La notación 5×3+ significa cinco series de tres repeticiones, donde el símbolo + indica que la última serie es AMRAP, es decir, tantas repeticiones como sea posible con técnica correcta. Si el usuario completa la sesión, se aplica add_weight y sube el peso. Si falla, se aplica advance_stage y pasa a la siguiente etapa.

Partiendo de un Training Max de 100 kg, la evolución simplificada sería:

- Sesión 1: prescribe 5×3+ a 85 kg. El usuario completa [3, 3, 3, 3, 5]. Resultado: éxito. Regla aplicada: add_weight(+5 kg). Nuevo Training Max: 105 kg.
- Sesión 2: prescribe 5×3+ a 90 kg. El usuario completa [3, 3, 3, 3, 4]. Resultado: éxito. Regla aplicada: add_weight(+5 kg). Nuevo Training Max: 110 kg.
- Sesión 3: prescribe 5×3+ a 95 kg. El usuario registra [3, 3, 2, 1, 0]. Resultado: fallo. Regla aplicada: advance_stage. La próxima sesión pasa a 6×2+.

Este ejemplo muestra que el estado visible para el usuario —peso, etapa y siguiente sesión— no se introduce manualmente. Se deriva del historial. También explica por qué deshacer una acción es sencillo: basta con eliminar la última entrada del historial y volver a calcular.

Algoritmo de replay
El cálculo del estado se implementa como un fold determinista sobre el historial ordenado. En términos no técnicos, un fold determinista significa recorrer todos los resultados anteriores en orden y aplicar siempre las mismas reglas. Si el historial y la definición son iguales, el resultado final también será siempre igual.

```typescript
function computeGenericProgram(definition, history, userMaxes) {
  let state = initialState(definition, userMaxes);

  for (const result of history) {
    const slot = state.slots[result.slotId];
    const stage = slot.definition.stages[slot.stageIndex];
    const rule = result.outcome === "success"
      ? stage.successRule
      : stage.failureRule;

    state = applyRule(state, slot, rule);
  }

  return state;
}
```

La complejidad temporal es O(n) respecto al número de resultados del historial. Para los volúmenes habituales del dominio —centenares de sesiones por programa, no millones de registros— el recálculo es suficientemente rápido en cliente y servidor. Esta decisión evita guardar estados redundantes y facilita la auditabilidad: cualquier peso prescrito puede explicarse reconstruyendo las sesiones que llevaron hasta él. Para la estimación de marcas y métricas derivadas se utiliza, entre otras, la fórmula de Epley, habitual en entrenamiento de fuerza (Epley, 1985).

Por qué este modelo encaja en Gravity Room
Un modelo mutable clásico permitiría guardar directamente “el peso que toca hoy” en una tabla, pero obligaría a mantener ese valor sincronizado cada vez que el usuario completa, edita, importa o deshace una sesión. En Gravity Room el cálculo es barato y el historial es pequeño, por lo que resulta más robusto recalcular el estado que persistirlo de forma redundante. La ventaja más importante es conceptual: la aplicación puede explicar de dónde sale cada prescripción, ejecutar el mismo motor en modo invitado y sincronizar dispositivos sin reglas especiales de reconciliación.
'''

hitos = '''5.2.11 Hitos de desarrollo
El proyecto se organizó en hitos incrementales, cada uno culminado con una versión funcional o un despliegue en producción:

1. Diseño del motor declarativo (computeGenericProgram): definición de la estructura de un programa —días, slots, etapas y reglas de progresión— y del algoritmo de replay.
2. API base y autenticación: Google OAuth, JWT con rotación dual y detección de reutilización de refresh tokens.
3. SPA con modo invitado: persistencia local validada con Zod e importación al servidor en el primer inicio de sesión.
4. Catálogo de programas: GZCLP como primer programa completo y ampliación progresiva hasta los ocho programas actuales.
5. PWA e internacionalización: instalación en dispositivos móviles y soporte completo en español e inglés.
6. Despliegue inicial en Railway: primera versión pública servida desde una plataforma Platform-as-a-Service para validar el producto bajo tráfico real.
7. Migración a VPS propio en Hetzner: Docker Compose, Caddy, GitHub Actions, health checks, validación previa de variables de entorno y pg_dump diario.
8. Observabilidad: logs estructurados con Pino y métricas Prometheus.
9. Catálogo abierto a la comunidad: separación entre program_definitions y program_templates, con flujo draft → pending_review → approved → rejected.

Figura 10. Diagrama de Gantt simplificado del desarrollo del proyecto. Elaboración propia.
'''

glossary = '''8. Anexos
8.1 Anexo I. Glosario de términos técnicos
Slot: posición concreta dentro de un día de entrenamiento. Por ejemplo, en un Día A puede haber un slot para sentadilla, otro para press banca y otro para remo.

AMRAP: siglas de “As Many Reps As Possible”. Indica que en la última serie el usuario realiza tantas repeticiones como pueda manteniendo buena técnica.

Etapa progresiva: fase de un ejercicio que define cuántas series y repeticiones debe realizar el usuario. Por ejemplo, 5×3+, 6×2+ y 10×1+ son tres etapas distintas.

5×3+: cinco series de tres repeticiones. El símbolo + indica que la última serie es AMRAP.

Training Max: peso de referencia usado por el programa para calcular las cargas de entrenamiento. No siempre coincide con el máximo real del usuario.

add_weight: regla que incrementa el peso si el usuario completa correctamente la sesión.

advance_stage: regla que pasa al usuario a la siguiente etapa cuando no consigue completar la sesión prescrita.

advance_stage_add_weight: regla que avanza de etapa y, además, incrementa el peso.

deload_percent: reducción controlada del peso, expresada como porcentaje, para facilitar la recuperación tras fallos acumulados.

add_weight_reset_stage: regla que aumenta el peso y vuelve a la primera etapa del ejercicio.

no_change: regla que mantiene la misma prescripción para la siguiente sesión.

Replay: proceso de reconstruir el estado actual aplicando una a una las entradas del historial.

Fold determinista: recorrido ordenado del historial en el que las mismas entradas producen siempre el mismo resultado final.

Event sourcing simplificado: enfoque en el que el estado se reconstruye a partir de eventos pasados, sin almacenar como dato principal el resultado final ya calculado.

JWT: token firmado utilizado para autenticar peticiones entre cliente y servidor.

Refresh token: token de larga duración usado para obtener nuevos access tokens sin pedir al usuario que inicie sesión de nuevo.

Reverse proxy: servidor que recibe las peticiones externas y las redirige internamente al servicio correspondiente, por ejemplo de api.gravityroom.app al contenedor de API.

CI/CD: integración y despliegue continuos. Automatización que compila, prueba y despliega el proyecto cuando se actualiza el repositorio.

8.2 Anexo II. Definición declarativa reducida del programa
La siguiente estructura resume cómo se representa un programa dentro del paquete compartido @gzclp/shared. Se incluye como anexo para no interrumpir el hilo principal de la memoria con un bloque de código excesivamente largo.

```typescript
export interface ProgramDefinition {
  id: string;
  name: string;
  level: "beginner" | "intermediate" | "advanced";
  category: "strength" | "hypertrophy";
  days: Day[];
}

export interface Day {
  dayIndex: number;
  label: string;
  slots: Slot[];
}

export interface Slot {
  slotId: string;
  exerciseId: string;
  stageIndex: number;
  stages: Stage[];
}

export interface Stage {
  sets: number;
  reps: number;
  amrap: boolean;
  intensityPct: number;
  successRule: ProgressionRule;
  failureRule: ProgressionRule;
}

export type ProgressionRule =
  | { type: "add_weight"; kg: number }
  | { type: "advance_stage" }
  | { type: "advance_stage_add_weight"; kg: number }
  | { type: "deload_percent"; pct: number }
  | { type: "add_weight_reset_stage"; kg: number }
  | { type: "no_change" };
```
'''

# Replace sections. Keep patterns anchored to body text so we do not accidentally
# match entries in the index.
text = re.sub(r'1\. Resumen\n.*?\n2\. Índice', summary + '\n2. Índice', text, flags=re.S)
text = re.sub(r'5\.2 Desarrollo del proyecto\n5\.2\.1 Elección del stack y justificación\n.*?\n5\.2\.2 Arquitectura general', '5.2 Desarrollo del proyecto\n' + stack + '\n5.2.2 Arquitectura general', text, flags=re.S)
text = re.sub(r'5\.2\.3 Motor de progresión declarativo\nEl\s+corazón.*?\n5\.2\.4 Despliegue e integración continua', motor + '\n5.2.4 Despliegue e integración continua', text, flags=re.S)
text = re.sub(r'5\.2\.11 Hitos de desarrollo\nEl proyecto se organizó.*?\n6\. Conclusiones', hitos + '\n6. Conclusiones', text, flags=re.S)
text = re.sub(r'2\. Índice\n.*?\n3\. Justificación', index + '\n3. Justificación', text, flags=re.S)
# The original extracted PDF index has line-number-only entries. Remove the leftover
# part of that old index so the real section 3 starts immediately after the new index.
text = re.sub(
    r'(8\.2 Anexo II\. Definición declarativa reducida del programa\n)\n3\. Justificación\n4\n.*?\n3\. Justificación\nEl sector',
    r'\1\n3. Justificación\nEl sector',
    text,
    flags=re.S,
)

# Targeted APA/cclarity fixes in unchanged sections
text = text.replace('normativa española (Ley\n31/1995 de Prevención de Riesgos Laborales, R.D. 488/1997 sobre puestos\ncon\npantallas\nde\nvisualización\nde\ndatos\ny\nnormativa\nautonómica\ncomplementaria)', 'normativa española (Jefatura del Estado, 1995; Ministerio de Trabajo y Asuntos Sociales, 1997)')
text = text.replace('A las medidas de prevención tradicionales se suma el derecho a la\ndesconexión digital, reconocido en la Ley Orgánica 3/2018 de Protección de\nDatos y Garantía de los Derechos Digitales y en la Ley 10/2021 de trabajo a\ndistancia.', 'A las medidas de prevención tradicionales se suma el derecho a la desconexión digital, reconocido en la Ley Orgánica 3/2018 de Protección de Datos y Garantía de los Derechos Digitales y en la Ley 10/2021 de trabajo a distancia (Jefatura del Estado, 2021).')
text = text.replace('PostgreSQL\n18 (cuyos datos viven', 'PostgreSQL 18 (PostgreSQL Global Development Group, 2025), cuyos datos viven')
text = text.replace('el\ntráfico en el edge antes de alcanzar el host (sólo 22, 80 y 443 abiertos), y el\nfirewall del host (ufw) repite la misma lista a nivel de sistema operativo.', 'el tráfico en el edge antes de alcanzar el host. Sólo se mantienen abiertos los puertos 22, 80 y 443: el puerto 22 para acceso SSH administrativo, el puerto 80 para tráfico HTTP y redirección a HTTPS, y el puerto 443 para tráfico HTTPS seguro. El firewall del host (ufw) repite la misma lista a nivel de sistema operativo.')
text = text.replace('Esta estrategia sigue las recomendaciones de OAuth 2.0 BCP4 y\nprotege contra ataques de reutilización de refresh tokens robados.', 'Esta estrategia sigue las recomendaciones actuales de OAuth 2.0 Security Best Current Practice (Bendell, Parecki et al., 2024) y protege contra ataques de reutilización de refresh tokens robados.')
text = text.replace('La autenticación se basa exclusivamente en Google OAuth:', 'La autenticación se basa exclusivamente en Google OAuth, siguiendo la documentación oficial de Sign in with Google for Web (Google, 2025):')
text = text.replace('métricas Prometheus\nen /metrics', 'métricas Prometheus (Prometheus Authors, 2024) en /metrics')
text = text.replace('Playwright sobre Chromium que validan los flujos críticos', 'Playwright sobre Chromium (Playwright Team, 2025) que validan los flujos críticos')
text = text.replace('Tailwind CSS 4 para los estilos.', 'Tailwind CSS 4 para los estilos (Tailwind Labs, 2025).')

# Add annex at end
text = text.rstrip() + '\n\n' + glossary

# Clean broken line wrapping into readable paragraphs while keeping code fences
lines = text.splitlines()
out=[]
in_code=False
buf=[]

def flush():
    global buf
    if buf:
        out.append(' '.join(x.strip() for x in buf if x.strip()))
        buf.clear()

heading_re=re.compile(r'^(\d+(?:\.\d+)*\.?\s+.+|[A-ZÁÉÍÓÚÑ][A-Za-zÁÉÍÓÚÜÑáéíóúüñ /+()→.-]{0,80})$')
for line in lines:
    s=line.rstrip()
    if s.startswith('```'):
        flush(); out.append(s); in_code=not in_code; continue
    if in_code:
        out.append(s); continue
    if not s.strip():
        flush(); out.append(''); continue
    # Keep bullets and numbered lists as separate lines
    if re.match(r'^[-•]\s+', s.strip()) or re.match(r'^\d+\.\s+', s.strip()):
        flush(); out.append(s.strip()); continue
    if re.match(r'^Figura \d+\.', s.strip()) or re.match(r'^Tabla \d+\.', s.strip()):
        flush(); out.append(s.strip()); continue
    if re.match(r'^\d+(?:\.\d+)*\s+', s.strip()):
        flush(); out.append(s.strip()); continue
    # code-ish lines outside fences
    if s.strip().startswith(('export ', 'function ', '|', '//', '{', '}', '  ', 'let ', 'for ', 'const ', 'return ')):
        flush(); out.append(s); continue
    buf.append(s)
flush()
md='\n'.join(out)

# Insert APA citations for bibliography refs if still absent in obvious terms
md = md.replace('Pino para logs estructurados y Prometheus para métricas.', 'Pino para logs estructurados y Prometheus para métricas (Prometheus Authors, 2024).')

md_path=ROOT/'RecheAmado_LuisLucas_Memoria_DAW_actualizada.md'
md_path.write_text(md, encoding='utf-8')

# Generate Gantt image
img_path=ROOT/'memoria_images'/'fig10_gantt.png'
img_path.parent.mkdir(exist_ok=True)
W,H=1600,720
img=Image.new('RGB',(W,H),'white')
d=ImageDraw.Draw(img)
try:
    font=ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',24)
    small=ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',20)
    bold=ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',24)
except:
    font=small=bold=None
left=360; top=80; row_h=54; col_w=115
months=['Nov','Dic','Ene','Feb','Mar','Abr','May','Jun']
for i,m in enumerate(months):
    x=left+i*col_w
    d.rectangle([x,30,x+col_w,top],outline=(180,180,180),fill=(245,245,245))
    d.text((x+35,48),m,fill=(0,0,0),font=bold)
items=[
('Diseño del motor',0,2),('API y autenticación',1,3),('SPA y modo invitado',2,4),('Catálogo de programas',3,6),('PWA e i18n',4,6),('Despliegue Railway',3,4),('Migración a Hetzner',5,6),('Observabilidad',5,7),('Catálogo comunidad',6,7),('Memoria y revisión',6,8)]
for r,(name,start,end) in enumerate(items):
    y=top+r*row_h
    d.text((20,y+15),name,fill=(0,0,0),font=small)
    for i in range(len(months)):
        x=left+i*col_w
        d.rectangle([x,y,x+col_w,y+row_h],outline=(220,220,220))
    d.rectangle([left+start*col_w+8,y+12,left+end*col_w-8,y+row_h-12],fill=(197,137,43),outline=(120,80,20))
d.text((20,15),'Cronograma simplificado del desarrollo de Gravity Room',fill=(0,0,0),font=bold)
img.save(img_path)

# Create docx
doc=Document()
sec=doc.sections[0]
sec.top_margin=Inches(0.8); sec.bottom_margin=Inches(0.8); sec.left_margin=Inches(0.9); sec.right_margin=Inches(0.9)
styles=doc.styles
styles['Normal'].font.name='Arial'; styles['Normal'].font.size=Pt(10.5)
styles['Heading 1'].font.name='Arial'; styles['Heading 1'].font.size=Pt(16)
styles['Heading 2'].font.name='Arial'; styles['Heading 2'].font.size=Pt(13)

fig_map={
 'Figura 1.': ROOT/'memoria_images/p14_img39.png',
 'Figura 2.': ROOT/'memoria_images/p19_img53.png',
 'Figura 3.': ROOT/'memoria_images/p21_img59.png',
 'Figura 4.': ROOT/'memoria_images/p23_img65.png',
 'Figura 5.': ROOT/'memoria_images/p29_img79.png',
 'Figura 6.': ROOT/'memoria_images/p30_img84.png',
 'Figura 7.': ROOT/'memoria_images/p31_img87.png',
 'Figura 8.': ROOT/'memoria_images/p32_img90.png',
 'Figura 10.': img_path,
}

in_code=False
for raw in md.splitlines():
    line=raw.strip()
    if not line:
        continue
    if line.startswith('```'):
        in_code=not in_code; continue
    if in_code:
        p=doc.add_paragraph()
        run=p.add_run(raw)
        run.font.name='Courier New'; run.font.size=Pt(8.5)
        continue
    # title page special
    if line=='GRAVITY ROOM':
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(line); r.bold=True; r.font.size=Pt(22)
        continue
    m=re.match(r'^(\d+)\.\s+(.+)$', line)
    m2=re.match(r'^(\d+\.\d+)\s+(.+)$', line)
    m3=re.match(r'^(\d+\.\d+\.\d+)\s+(.+)$', line)
    if m3:
        doc.add_heading(line, level=3); continue
    if m2:
        doc.add_heading(line, level=2); continue
    if m and int(m.group(1)) in range(1,9):
        doc.add_heading(line, level=1); continue
    if any(line.startswith(k) for k in fig_map):
        path=fig_map[next(k for k in fig_map if line.startswith(k))]
        if path.exists():
            try:
                doc.add_picture(str(path), width=Inches(6.5))
                doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass
        p=doc.add_paragraph(line)
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.italic=True; r.font.size=Pt(9)
        continue
    if line.startswith('- '):
        doc.add_paragraph(line[2:], style='List Bullet'); continue
    if re.match(r'^\d+\.\s+', line):
        doc.add_paragraph(re.sub(r'^\d+\.\s+','',line), style='List Number'); continue
    p=doc.add_paragraph(line)

out_docx=ROOT/'RecheAmado_LuisLucas_Memoria_DAW_actualizada.docx'
doc.save(out_docx)
print(md_path)
print(out_docx)
print(img_path)
